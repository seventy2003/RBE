# -*- coding: utf-8 -*-

"""
FILE : view 

DESCRIPTION:   
                                                                                       
FUNCTION LIST:  
1. UI
2. get input
3. debug 

REVISION:  
  
Ver  | yyyymmdd | Who    | Description of changes
1.00 | 20181105 | wxhao  | Create.
1.00 | 20191121 | wxhao  | Add ' encoding="utf-8-sig" ' to support chinese character 
     in configure file

"""

import sys
import os
import configparser
from docx import Document
import csv

from strace import *
from req import *
from inform import *
from parse import Parse
from output import OutputRBE
from relation import Relation
from critical import Critical
from verify import Verify


# for debug display
from req import srsDict  

# for debug
import time

cmdUI = [
        '--------------------------------'  ,
        'Input command:'                    ,
        '  \'1\': start parse'              ,
        '  \'2\': display list'             ,
        '  \'3\': output SRS'               ,
        '  \'4\': trace USER ID'            ,
        '  \'5\': relation analyze'         ,
        '  \'6\': critical analyze'         ,
        '  \'7\': verify table'             ,        
        '  \'8\': tradeoff analyze'         ,
        '  \'9\': generate figures'         ,                    
        '  \'0\': exit.'
]

subUIFig = [
        '--------------------------------'  ,
        'Input command:'                    ,
        '  \'1\': SRS Pie chart'            ,
        '  \'2\': verify Pie chart'         ,                          
        '  \'0\': return.'
]

subUITrace = [
        '--------------------------------'  ,
        'Input command:'                    ,
        '  \'1\': SRS and USER'             ,
        '  \'2\': SDD and SRS'              ,                          
        '  \'0\': return.'
]

class View:
    
    def __init__(self):
        pass
    def init(self):
        pass
    def getInput(self):
        pass    
    def interact(self):
        pass
    def run(self):
        pass

class CmdView(View):

    parse = None
    trace = None
    outPutSrs = None
    relat = None
    crt = None
    verif = None
    infm = None
    
    def init(self):
        self.parse = Parse()
        self.outPutSrs = OutputRBE()
        self.relat = Relation()
        self.crt = Critical()
        self.verif = Verify()
        self.infm = Inform()

    def getInput(self):

        #self.parse = Parse()

        # get srs file
        try:
            cf = configparser.ConfigParser()
            cf.read("smg.conf",encoding="utf-8-sig")
            infile = cf.get("input", "inSrsFile")
            #self.f = open(infile, 'r', encoding = 'gb18030', errors = 'ignore')
            self.f = Document(infile)
        except IOError:
            print("File Open fail\n")
            raise Exception("File Open fail\n")

        #print('Done.')

    def interact(self):

        # display UI
        for line in cmdUI:
            print(line)

        cmd = input()
        if cmd == '1':
            #clear for next trace
            self.parse.state = 0
            srsDict.clear()
            if self.trace:
                self.trace.tClear()
            
            print("Start parse...")

            # for performance watch
            startTime = time.time()

            self.parse.doParse(self.f)
            #self.parse.run(self.f)

            # for performance watch
            endTime = time.time()

            print('Done.\n')

            # for performance watch
            print("Time is %s Second\n" %(endTime - startTime))
            
        elif cmd == '2':
            self.display()
            #self.infm.checkAll()
        elif cmd == '3':
            print('Output SRS...')
            if self.parse.state == 100:
                self.outPutSrs.srsOutput()
                print("Done.")
            else:
                print('Parse first please.')

        elif cmd == '4':
            if self.parse.state == 100:
                self.subTraceInteract()
            else:
                print('Parse first please.')
        elif cmd == '5':
            print('Analyze relations of SRS...')
            if self.parse.state == 100:
                self.relat.genLevelRt()
                self.relat.outputRt()
                print("Done.")
            else:
                print('Parse first please.')
        elif cmd == '6':
            print('Analyze relations of SRS...')
            if self.parse.state == 100:
                self.crt.outputCrtal()
                print("Done.")
            else:
                print('Parse first please.')
        elif cmd == '7':
            print('Generate verify table...')
            if self.parse.state == 100:
                self.verif.outputVerify()
                print("Done.")
            else:
                print('Parse first please.')
        elif cmd == '9':
            self.subFigInteract()

        elif cmd == '0':
            os._exit(0)
        else:
            print('Invalid input.')
            
        return

    def subTraceInteract(self):
        
        while (1):
            # display UI
            for line in subUITrace:
                print(line)

            cmd = input()
            if cmd == '1':
                print('SRS and USR tracing.')
                # get rw file
                try:
                    cf = configparser.ConfigParser()
                    cf.read("smg.conf",encoding="utf-8-sig")
                    infile = cf.get("input", "inUsrFile")
                    #self.f = open(infile, 'r', encoding = 'gb18030', errors = 'ignore')
                    self.fileRw = Document(infile)
                except IOError:
                    print("File Open fail\n")
                    raise Exception("File Open fail\n")

                self.trace = STrace(self.fileRw)
                self.trace.getUsrDict()
                self.trace.doTraceSrs2Usr()
                self.trace.debug()

            elif cmd == '2':
                print('SDD and SRS tracing.')

                # get sdd file
                try:
                    cf = configparser.ConfigParser()
                    cf.read("smg.conf",encoding="utf-8-sig")
                    infile = cf.get("input", "inSddFile")
                    #self.f = open(infile, 'r', encoding = 'gb18030', errors = 'ignore')
                    self.fileRw = Document(infile)
                except IOError:
                    print("File Open fail\n")
                    raise Exception("File Open fail\n")

                self.trace = STrace(self.fileRw)
                self.trace.getSddTraceList()
                self.trace.doTraceSdd2Srs()
                self.trace.sddTrOutput()

            elif cmd == '0':
                return
            else:
                print('Invalid input.')   
    
    def subFigInteract(self):

        while (1):
            # display UI
            for line in subUIFig:
                print(line)

            cmd = input()
            if cmd == '1':
                print('Print SRS pie chart.')
                self.outPutSrs.genPieChart()
            elif cmd == '2':
                print('Print verify pie chart.')
                self.verif.genPieChart()
            elif cmd == '0':
                return
            else:
                print('Invalid input.')      
        
    def run(self):
        self.init()
        while (1):
            self.getInput()
            self.interact()

    # for debug
    def display(self):
        # for debug
        for key in srsDict:
            if srsDict[key].type == 'FUNC':
                print(srsDict[key].id, "    ", srsDict[key].dataIn, "    ",
                      srsDict[key].dataOut, "    ", srsDict[key].trace, "    ",
                      srsDict[key].verify)
            if srsDict[key].type == 'INTF':
                print(srsDict[key].id, "    ", srsDict[key].trace, "    ",
                      srsDict[key].verify)

        print("Dsplay for debug Done.")

        



